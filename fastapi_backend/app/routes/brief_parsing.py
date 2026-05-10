"""
Brief parsing endpoint: CA uploads an Excel or document file →
items are extracted and created as VerificationItems for the audit.
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import List
from app.core.database import get_db
from app.models.verification_item import VerificationItem, VerificationStatus, ItemType
from app.models.document import Document, DocumentCategory, DocumentType
from app.models.user import User
from app.routes.auth import get_current_user
from app.models.audit import Audit
from app.services.ai_service import parse_brief_with_gemini
from app.utils.s3 import upload_file_to_s3
from app.schemas.verification_item import VerificationItemResponse
from app.schemas.document import DocumentResponse
from datetime import datetime
import io
import logging

logger = logging.getLogger(__name__)

router = APIRouter(tags=["brief_parsing"])

EXCEL_EXTENSIONS = {".xlsx", ".xls"}
TEXT_EXTRACTABLE = {".pdf", ".docx", ".doc", ".txt", ".csv"}


def _extract_text_from_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_from_docx(content: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


def _parse_excel(content: bytes) -> List[dict]:
    """Parse Excel rows into verification item dicts.

    Expected columns (case-insensitive, order flexible):
    type | title/name/description | reference_value/value/reference
    Falls back gracefully when columns are missing.
    """
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    # Detect header row
    headers = [str(c).strip().lower() if c else "" for c in rows[0]]
    data_rows = rows[1:] if any(h.isalpha() for h in headers) else rows

    ITEM_TYPE_MAP = {
        "vehicle": ItemType.VEHICLE, "car": ItemType.VEHICLE, "truck": ItemType.VEHICLE,
        "property": ItemType.PROPERTY, "land": ItemType.PROPERTY, "building": ItemType.PROPERTY,
        "equipment": ItemType.EQUIPMENT, "machinery": ItemType.EQUIPMENT,
        "inventory": ItemType.INVENTORY, "stock": ItemType.INVENTORY,
        "bank": ItemType.BANK_ACCOUNT, "account": ItemType.BANK_ACCOUNT,
        "financial": ItemType.FINANCIAL_RECORD, "record": ItemType.FINANCIAL_RECORD,
    }

    def _col(row, *names):
        for name in names:
            for i, h in enumerate(headers):
                if name in h and i < len(row):
                    val = row[i]
                    return str(val).strip() if val is not None else ""
        # fallback: return first non-empty cell
        for cell in row:
            if cell is not None:
                return str(cell).strip()
        return ""

    def _item_type(raw: str) -> str:
        raw = raw.lower().strip()
        for key, val in ITEM_TYPE_MAP.items():
            if key in raw:
                return val
        return ItemType.OTHER

    items = []
    for row in data_rows:
        if all(c is None for c in row):
            continue
        title_raw = _col(row, "title", "name", "item", "description", "asset")
        type_raw = _col(row, "type", "category", "kind")
        ref_raw = _col(row, "reference", "ref", "value", "reg", "id", "number", "amount")
        desc_raw = _col(row, "detail", "note", "remark", "description")

        if not title_raw:
            continue

        items.append({
            "title": title_raw,
            "item_type": _item_type(type_raw),
            "description": desc_raw,
            "reference_value": ref_raw,
            "is_ai_parsed": False,
        })

    return items


@router.post("/audits/{audit_id}/parse-brief")
async def parse_brief(
    audit_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Upload an audit brief and auto-create verification items from it."""
    audit = db.query(Audit).filter(
        Audit.id == audit_id,
        Audit.workspace_id == current_user.workspace_id,
    ).first()
    if not audit:
        raise HTTPException(status_code=404, detail="Audit not found")

    content = await file.read()
    filename = file.filename or "brief"
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Save brief to S3
    s3_key = f"audits/{audit_id}/briefs/{datetime.utcnow().timestamp()}_{filename}"
    try:
        file_url = await upload_file_to_s3(content, s3_key)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to store brief: {e}")

    brief_doc = Document(
        audit_id=audit_id,
        uploaded_by=current_user.id,
        file_name=filename,
        file_path=s3_key,
        file_size=len(content),
        document_category=DocumentCategory.BRIEF,
        document_type=DocumentType.OTHER,
        workspace_id=current_user.workspace_id,
    )
    db.add(brief_doc)
    db.flush()

    # Extract verification items
    if ext in EXCEL_EXTENSIONS:
        raw_items = _parse_excel(content)
    else:
        # Extract text for Gemini
        try:
            if ext == ".pdf":
                text = _extract_text_from_pdf(content)
            elif ext in (".docx", ".doc"):
                text = _extract_text_from_docx(content)
            else:
                text = content.decode("utf-8", errors="replace")
        except Exception as e:
            logger.error("Text extraction failed: %s", e)
            text = ""

        if not text.strip():
            db.commit()
            return {"brief_document": brief_doc.id, "items_created": [], "warning": "Could not extract text from document"}

        raw_items = await parse_brief_with_gemini(text)
        for item in raw_items:
            item["is_ai_parsed"] = True

    created = []
    for raw in raw_items:
        item = VerificationItem(
            audit_id=audit_id,
            title=raw["title"],
            description=raw.get("description") or None,
            item_type=raw.get("item_type", ItemType.OTHER),
            reference_value=raw.get("reference_value") or None,
            is_ai_parsed=raw.get("is_ai_parsed", False),
            workspace_id=current_user.workspace_id,
        )
        db.add(item)
        created.append(item)

    db.commit()
    for item in created:
        db.refresh(item)

    return {
        "brief_document_id": brief_doc.id,
        "items_created": [
            {
                "id": item.id,
                "title": item.title,
                "item_type": item.item_type,
                "reference_value": item.reference_value,
                "is_ai_parsed": item.is_ai_parsed,
            }
            for item in created
        ],
    }
